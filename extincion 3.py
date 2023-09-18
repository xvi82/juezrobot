import tkinter as tk
from tkinter import messagebox, ttk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkcalendar import DateEntry
from tkinter.scrolledtext import ScrolledText
import openai
import threading
import math
from datetime import datetime, timedelta
import re
import babel.numbers

# Configura la calculadora indemnizatoria

def indemnizacion_despido(fecha_antiguedad, salari, fecha_despido):
    fechalimite = datetime.strptime('12/02/2012', '%d/%m/%Y').date()
    antig = datetime.strptime(fecha_antiguedad, '%d/%m/%Y').date()
    despid = datetime.strptime(fecha_despido, '%d/%m/%Y').date()
    salario = float(re.sub(',', '.', salari))
    if antig > fechalimite:
        dif = math.ceil(((despid - antig) / timedelta(days=1))) + 1
        daf = math.ceil(dif / 30.41666667)
        indem = daf * salario * 2.75
        return indem
    else:
        dif1 = math.ceil(((fechalimite - antig) / timedelta(days=1))) + 1
        daf1 = math.ceil((dif1 / 30.41666667))
        indemprev = daf1 * salario * 3.75
        dif2 = math.ceil(((despid - fechalimite) / timedelta(days=1))) + 1
        daf2 = math.ceil(dif2 / 30.41666667)
        indempost = daf2 * salario * 2.75
        indem2 = indemprev + indempost
        return indem2

def calcular_indemnizacion():
    fecha_antiguedad = antiguedad_date_entry.get()
    salari = salario_entry.get()
    fecha_despido = datetime.today().strftime('%d/%m/%Y')

    return round(indemnizacion_despido(fecha_antiguedad, salari, fecha_despido), 2)

# Configura tu clave de API
openai.api_key = 'CLAVE API AQUI'

# Leer las reglas para resolver el pleito
with open('resolutor.txt', 'r', encoding="utf8") as file:
    reglas = file.read()


#crear hilo
def thread_safe_resolver():
    btn_resol.config(state=tk.DISABLED)  # Deshabilita el botón mientras se ejecuta la función
    mostrar_respuesta()
    btn_resol.config(state=tk.NORMAL)  # Habilita el botón una vez que finaliza la función

# Obtener la respuesta de GPT-4 para resolver el pleito
def obtener_respuesta(prompt):
    # Hacer la petición a GPT-4
    calcular_indemnizacion()
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": reglas},
            {"role": "user", "content": prompt},
        ]
    )
    texto = response['choices'][0]['message']['content'].strip()

    # Dividir el texto en frases (usando el punto como delimitador)
    frases = texto.split(".")

    # Inicializar variables
    respuesta_modificada = ""
    numero = None

    # Revisar cada frase para encontrar "CONDENAR" o "ABSOLVER"
    for frase in frases:
        if "CONDENAR" in frase:
            numero = 0
        elif "ABSOLVER" in frase:
            numero = 1
        else:
            respuesta_modificada += frase + "."

    # Limpieza final de la respuesta (eliminar el punto final)
    respuesta_modificada = respuesta_modificada.rstrip(".")

    return respuesta_modificada, numero

# Borra saltos de linea para GPT-4
def borrar_saltos_linea_vacios(texto):
        return texto.replace("\n\n", "\n")

# Mostrar la respuesta de GPT-4 en la GUI
# Mostrar la respuesta de GPT-4 en la GUI
def mostrar_respuesta():
    indemnizacion = calcular_indemnizacion()
    global modelo  # Hacemos que modelo sea global para poder modificarla dentro de esta función

    prompt_usuario = scrolled_text.get("1.0", tk.END).strip()
    respuesta, numero = obtener_respuesta(prompt_usuario)

    # Obtener el nombre del demandante y del demandado del cuadro de texto correspondiente
    nombre_demandante = entries["Nombre del demandante"].get()
    nombre_demandado = entries["Nombre de la institución demandada"].get()

    # Reemplazar [argumentación por la IA] con la respuesta de GPT-4
    modelo_actualizado = modelo.replace("[argumentación por la IA]", borrar_saltos_linea_vacios(respuesta))

    # Reemplazar [fallo] según el número
    if numero == 0:
        fallo_texto = f"""
QUINTO.- En cuanto a las consecuencias jurídicas de la declaración de extinción de la relación laboral, las mismas vienen contempladas en el artículo 56.1 del Estatuto de los Trabajadores. El empresario, en el plazo de cinco días desde la notificación de la sentencia, puede optar entre la readmisión del actor con abono de los salarios de tramitación, y el pago de la indemnización prevista en el artículo 56.1.a) del Estatuto de los Trabajadores; esto es, una indemnización igual al importe de treinta y tres días de salario por años trabajado con prorrateo de los periodos inferiores al años y con un máximo de cuarenta y dos mensualidades, a lo que deberá añadirse también el importe de los salarios de tramitación en la forma prevista en el artículo 56.1.b) del Estatuto de los Trabajadores.

SEXTO.-  Habiéndose citado al Fondo de Garantía Salarial, conforme a lo dispuesto en el artículo 23.1 de la Ley Reguladora de la Jurisdicción Social, únicamente puede ser condenado a estar y pasar por este pronunciamiento, sin perjuicio de resultar ulteriormente la insolvencia de la empresa, deba asumir su responsabilidad legal.

SÉPTIMO.- En virtud de lo dispuesto en el art. 191.3.a) de la Ley Reguladora de la Jurisdicción Social, contra esta Sentencia puede interponerse Recurso de Suplicación

     Vistos los preceptos legales citados y demás de general observancia y por la autoridad que me confiere el art. 117 de la Constitución Española y 1 de la Ley Orgánica del Poder judicial,

FALLO

Que DEBO ESTIMAR y ESTIMO la demanda interpuesta por {nombre_demandante} contra {nombre_demandado} y FOGASA, y por ende debo declarar y declaro extinguida la relación laboral que le unía a , con efectos desde esta resolución judicial, condenando a la empresa a abonarle en concepto de indemnización la cantidad de {indemnizacion} euros.

Notifíquese la presente Resolución a las partes en legal forma, haciéndose saber al tiempo que contra la misma cabe recurso de Suplicación, para ante la Sala de lo Social del Tribunal Superior de Justicia de Canarias, y que deberá anunciarse ante este Juzgado en el plazo de 5 días hábiles siguientes a la notificación de esta sentencia, siendo indispensable que el recurrente que no goce del beneficio de justicia gratuita acredite, al anunciar el recurso, haber consignado la cantidad objeto de la condena, que podrá sustituirse por el aseguramiento mediante aval bancario. La consignación deberá efectuarse en la Cuenta de Depósitos y Consignaciones de este Juzgado.

Así, por ésta, mi Sentencia, de la que se llevará certificación a los autos, lo pronuncio, mando y firmo."""

    else:
        fallo_texto = f"""
QUINTO.- En virtud de lo dispuesto en el art. 191.3.a) de la Ley Reguladora de la Jurisdicción Social, contra esta Sentencia puede interponerse Recurso de Suplicación

     Vistos los preceptos legales citados y demás de general observancia y por la autoridad que me confiere el art. 117 de la Constitución Española y 1 de la Ley Orgánica del Poder judicial,

FALLO

Que DEBO DESESTIMAR y DESESTIMO la demanda interpuesta por {nombre_demandante} contra {nombre_demandado} y FOGASA, y por ende absuelvo a la demandada de todos los pedimentos efectuados en su contra.

Notifíquese la presente Resolución a las partes en legal forma, haciéndose saber al tiempo que contra la misma cabe recurso de Suplicación, para ante la Sala de lo Social del Tribunal Superior de Justicia de Canarias, y que deberá anunciarse ante este Juzgado en el plazo de 5 días hábiles siguientes a la notificación de esta sentencia, siendo indispensable que el recurrente que no goce del beneficio de justicia gratuita acredite, al anunciar el recurso, haber consignado la cantidad objeto de la condena, que podrá sustituirse por el aseguramiento mediante aval bancario. La consignación deberá efectuarse en la Cuenta de Depósitos y Consignaciones de este Juzgado.

Así, por ésta, mi Sentencia, de la que se llevará certificación a los autos, lo pronuncio, mando y firmo."""

    modelo_actualizado = modelo_actualizado.replace("[fallo]", fallo_texto)

    vista_previa.delete(1.0, tk.END)
    vista_previa.insert(tk.END, modelo_actualizado)

    modelo = modelo_actualizado  # Actualizamos la variable global modelo con el contenido modificado


def cargar_modelo(filename):
    with open(filename, 'r', encoding='utf-8') as file:
        return file.read()


def formatear_fecha(date_obj):
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    dia = date_obj.day
    mes = meses[date_obj.month - 1]
    año = date_obj.year
    return f"{dia} de {mes} de {año}"

# Definición de las variables en un ámbito global
texto_original = ("SEGUNDO.- Admitida a trámite la demanda, las partes fueron citadas al acto de juicio. "
                  "Comparecidas las partes, asistidas en la forma que consta en acta, se pasó al acto de juicio. "
                  "En él, y una vez que se hubo efectuado la dación de cuenta de los antecedentes, la parte actora "
                  "se ratificó en su demanda. La parte demandada se opuso a la demanda y la contestó formulando "
                  "las alegaciones que constan en acta. Seguidamente, fue abierta la fase probatoria, en la que "
                  "se practicaron las pruebas que, propuestas por las partes, fueron declaradas pertinentes y constan "
                  "documentadas en autos. Practicada la prueba, las partes informaron sobre sus pretensiones y el juicio "
                  "quedó visto para sentencia.")

texto_si = texto_original

texto_no = ("SEGUNDO.- Admitida a trámite la demanda, las partes fueron citadas al acto de juicio. "
            "Comparecida únicamente la parte actora, asistida en la forma que consta en acta, se pasó al acto de juicio. "
            "En él, y una vez que se hubo efectuado la dación de cuenta de los antecedentes, la parte actora "
            "se ratificó en su demanda. Seguidamente, fue abierta la fase probatoria, en la que se practicaron "
            "las pruebas que, propuestas por la parte actora, única compareciente, fueron declaradas pertinentes y constan "
            "documentadas en autos. Practicada la prueba, la parte informó sobre sus pretensiones y el juicio quedó visto "
            "para sentencia.")


def actualizar_vista(event=None):
    modelo_temp = modelo[:]

    # Actualizar modelo_temp con las entradas de los widgets (Fecha, Número de juicio, etc.)
    for key, widget in entries.items():
        if key == "Fecha":
            fecha_seleccionada = widget.get_date()
            fecha_formateada = formatear_fecha(fecha_seleccionada)
            modelo_temp = modelo_temp.replace(f"[{key}]", fecha_formateada)
        else:
            modelo_temp = modelo_temp.replace(f"[{key}]", widget.get())

    # Lógica del radiobutton
    if entries["Opciones"].get() == "Si":
        modelo_temp = modelo_temp.replace(texto_original, texto_si)
    else:
        modelo_temp = modelo_temp.replace(texto_original, texto_no)

    # Construir los hechos probados escritos
    prefixes = ["PRIMERO", "SEGUNDO", "TERCERO", "CUARTO", "QUINTO", "SEXTO", "SÉPTIMO", "OCTAVO", "NOVENO", "DÉCIMO"]
    hechos_probados_escritos = "\n\n".join([prefixes[i] + ".- " + textbox.get() for i, textbox in enumerate(textboxes) if textbox.get()])

    # Determinar la continuación de la numeración para los hechos probados en el modelo
    num_hechos_probados_escritos = len([textbox for textbox in textboxes if textbox.get()])
    hechos_probados_modelo = modelo_temp.split("HECHOS PROBADOS\n")[1].split("FUNDAMENTOS DE DERECHO")[0].strip().split("\n")

    # Eliminamos la numeración anterior en el modelo
    hechos_probados_modelo_limpio = [(hecho.split(".-")[1] if ".-" in hecho else hecho) for hecho in hechos_probados_modelo if hecho.strip()]

    # Añadimos la nueva numeración
    hechos_probados_modelo_numerados = [prefixes[num_hechos_probados_escritos + i] + ".-" + hecho for i, hecho in enumerate(hechos_probados_modelo_limpio)]

    # Construir el texto final de hechos probados
    hechos_probados_final = hechos_probados_escritos + "\n\n" + "\n\n".join(hechos_probados_modelo_numerados)

    # Insertar los hechos probados final en el modelo
    modelo_temp = modelo_temp.split("HECHOS PROBADOS\n")[0] + "HECHOS PROBADOS\n\n" + hechos_probados_final + "\n\nFUNDAMENTOS DE DERECHO" + modelo_temp.split("FUNDAMENTOS DE DERECHO")[1]

    # Lógica para los checkboxes de pruebas
    pruebas_seleccionadas_list = [pruebas_mapping[prueba] for prueba, var in pruebas_checkboxes.items() if var.get()]
    if len(pruebas_seleccionadas_list) > 1:
        last_prueba = pruebas_seleccionadas_list.pop()
        pruebas_seleccionadas = ", ".join(pruebas_seleccionadas_list) + " y " + last_prueba
    else:
        pruebas_seleccionadas = ", ".join(pruebas_seleccionadas_list)

    modelo_temp = modelo_temp.replace("[pruebas]", pruebas_seleccionadas)

    vista_previa.delete(1.0, tk.END)
    vista_previa.insert(tk.END, modelo_temp)


def combined_save_to_docx():
    calcular_indemnizacion()
    actualizar_vista()
    """Function to save the content from both the main GUI and the textboxes to a .docx file."""
    doc = Document()
    contenido = vista_previa.get(1.0, tk.END).split("\n")

    es_respuesta_inicio = False
    es_respuesta_final = False

    for linea in contenido:
        if "CUARTO.-" in linea:
            es_respuesta_inicio = True
            _agregar_linea_justificada_con_negrita(doc, linea)
        elif "QUINTO.-" in linea:
            es_respuesta_final = True
            _agregar_linea_justificada_con_negrita(doc, linea)
        elif es_respuesta_inicio and not es_respuesta_final:
            _agregar_linea_justificada(doc, linea)
        elif linea in ['SENTENCIA', 'ANTECEDENTES DE HECHO', 'FUNDAMENTOS DE DERECHO', 'FALLO']:
            _agregar_linea_centralizada_negrita(doc, linea)
        elif linea.startswith("TERCERO.-") and "autos se han observado" in linea:
            _agregar_linea_justificada_con_negrita(doc, linea)
        elif linea.startswith(('PRIMERO.-', 'SEGUNDO.-', 'TERCERO.-', 'CUARTO.-', 'QUINTO.-', 'SEXTO.-', 'SÉPTIMO', 'OCTAVO.-', 'NOVENO.-', 'DÉCIMO.-')):
            _agregar_linea_justificada_con_negrita(doc, linea)
        elif linea.startswith(("Visto por mí", 'Que DEBO', 'Vistos los preceptos legales', 'Notifíquese la presente', 'Así, por ésta, mi Sentencia')):
            _agregar_linea_justificada(doc, linea)
        elif linea == "HECHOS PROBADOS":
            _agregar_linea_centralizada_negrita(doc, linea)
        else:
            # Aquí es donde se agregan las líneas en blanco con el formato deseado
            par = doc.add_paragraph(linea)
            par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            par.paragraph_format.line_spacing = 1.0  # Espaciado entre líneas a 1,0
            par.paragraph_format.space_after = Pt(0)  # Espaciado posterior a 0 pt

    # El resto del código de la función para guardar el documento como .docx
    numero_procedimiento, año_procedimiento = entries["Número de juicio"].get().split("/")
    nombre_archivo = f"JO_extincion_{numero_procedimiento}-{año_procedimiento}.docx"
    doc.save(nombre_archivo)
    messagebox.showinfo("Guardado", f"Documento guardado como {nombre_archivo}")


def _agregar_linea_centralizada_negrita(doc, linea):
    par = doc.add_paragraph()
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = par.add_run(linea)
    run.bold = True
    par.paragraph_format.line_spacing = 1.0  # Espaciado entre líneas a 1,0
    par.paragraph_format.space_after = Pt(0)  # Espaciado posterior a 0 pt

def _agregar_linea_justificada_con_negrita(doc, linea):
    par = doc.add_paragraph()
    par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    start_word, rest_of_line = linea.split(" ", 1)
    run = par.add_run(start_word)
    run.bold = True
    par.add_run(f" {rest_of_line}")
    par.paragraph_format.line_spacing = 1.0  # Espaciado entre líneas a 1,0
    par.paragraph_format.space_after = Pt(0)  # Espaciado posterior a 0 pt


def _agregar_linea_justificada(doc, linea):
    par = doc.add_paragraph(linea)
    par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    par.paragraph_format.line_spacing = 1.0  # Espaciado entre líneas a 1,0
    par.paragraph_format.space_after = Pt(0)  # Espaciado posterior a 0 pt


def _crear_frame(parent):
    frame = tk.Frame(parent, relief="flat", borderwidth=1, bg="white")
    return frame

def add_textbox():
    """Función para añadir un nuevo cuadro de texto al UI."""
    textbox = tk.Entry(hechos_probados_frame, bg="white", fg="black", relief="sunken",
                       highlightthickness=1, font=FUENTE_GLOBAL)
    textbox.pack(fill=tk.X, pady=2, padx=10) # Asegurarse de que ocupe todo el ancho con fill=tk.X
    textbox.bind("<KeyRelease>", actualizar_vista)
    textboxes.append(textbox)

def remove_last_textbox():
    if textboxes:
        textbox = textboxes.pop()
        textbox.destroy()
        actualizar_vista()

def _crear_entradas(frame, keys):
    entries = {}
    for key in keys:
        label = tk.Label(frame, text=key, bg="white", font=FUENTE_GLOBAL)
        label.pack(anchor="w", padx=0, pady=0)

        if key == "Fecha":
            widget = DateEntry(frame, width=30, date_pattern='d/m/y', background="white", foreground="black",
                               relief="groove", highlightthickness=1)
            widget.configure(font=FUENTE_GLOBAL)
            widget.bind("<<DateEntrySelected>>", actualizar_vista)
        else:
            widget = tk.Entry(frame, width=30, bg="white", fg="black", relief="sunken", highlightthickness=1,
                              font=FUENTE_GLOBAL)
            widget.bind("<KeyRelease>", actualizar_vista)

        widget.pack(fill=tk.X, pady=2, padx=10)
        entries[key] = widget

    opciones_label = tk.Label(frame, text="¿Comparecen ambas partes?", bg="white", font=FUENTE_GLOBAL)
    opciones_label.pack(anchor="w", padx=0, pady=0)

    opciones_var = tk.StringVar(value="Si")
    rb_si = tk.Radiobutton(frame, text="Si", variable=opciones_var, value="Si", bg="white", font=FUENTE_GLOBAL,
                           command=actualizar_vista)
    rb_no = tk.Radiobutton(frame, text="No", variable=opciones_var, value="No", bg="white", font=FUENTE_GLOBAL,
                           command=actualizar_vista)
    rb_si.pack(anchor="w", padx=10, pady=0)
    rb_no.pack(anchor="w", padx=10, pady=0)

    entries["Opciones"] = opciones_var
    return entries

# Configuración inicial
modelo = cargar_modelo("extincion.txt")  # Asegúrate de cambiar esta ruta al lugar correcto en tu sistema
keys = ['Fecha', 'Número de juicio', 'Nombre del demandante', 'Nombre de la institución demandada']

FUENTE_GLOBAL = ("Segoe UI", 9)  # Definición de la fuente global

# Configuración de la GUI
app = tk.Tk()
app.title("Rellenador de Modelo")
app.configure(bg="white")

# Establecer el tamaño de la ventana
app.geometry("1200x750")

# Centrar la ventana en la pantalla
screen_width = app.winfo_screenwidth()
screen_height = app.winfo_screenheight()
window_width = 1200
window_height = 750
x_position = (screen_width - window_width) / 2
y_position = (screen_height - window_height) / 2
app.geometry(f"{window_width}x{window_height}+{int(x_position)}+{int(y_position)}")

frame_izq = _crear_frame(app)
frame_der = _crear_frame(app)

# Diccionario para almacenar las variables de los checkboxes de pruebas
pruebas_checkboxes = {}

# Lista de pruebas disponibles
pruebas_opciones = ["Documental", "Testifical", "Pericial", "Forense", "Interrogatorio dte", "Interrogatorio ddo"]
pruebas_mapping = {
    "Documental": "la documental aportada",
    "Testifical": "la testifical",
    "Pericial": "la pericial",
    "Forense": "la pericial forense",
    "Interrogatorio dte": "el interrogatorio de la parte demandante",
    "Interrogatorio ddo": "el interrogatorio de la parte demandada"
}

# Configura la distribución de los frames usando grid
frame_izq.grid(row=0, column=0, sticky="nsew", padx=3, pady=3)  # Agregar padding a los frames
frame_der.grid(row=0, column=1, sticky="nsew", padx=3, pady=3)  # Agregar padding a los frames

# Configura cómo se distribuye el espacio entre las columnas
app.grid_columnconfigure(0, weight=1)  # 1/4 para el frame izquierdo
app.grid_columnconfigure(1, weight=5)  # 3/4 para el frame derecho
app.grid_rowconfigure(0, weight=1)

entries = _crear_entradas(frame_izq, keys)

# Add the label 'Hechos Probados' and the new textboxes to the existing GUI
hechos_probados_label = tk.Label(frame_izq, text="Hechos Probados", bg="white", font=FUENTE_GLOBAL)
hechos_probados_label.pack(anchor="w", padx=0, pady=2) # Añadido padding para consistencia

hechos_probados_frame = tk.Frame(frame_izq, bg="white")
hechos_probados_frame.pack(fill=tk.X, padx=0, pady=2) # Añadido fill=tk.X y padding para consistencia

# List to keep reference to all the textboxes
textboxes = []

# Add the first textbox
add_textbox()

# Frame para contener los botones
btn_frame = tk.Frame(frame_izq, bg="white")
btn_frame.pack(pady=5)

# Botón de adición
btn_add = ttk.Button(btn_frame, text="+", command=add_textbox, width=3)
btn_add.pack(side="left", padx=2)  # Se usa pack con side="left" para que estén en la misma línea

# Botón de eliminación
btn_remove = ttk.Button(btn_frame, text="-", command=remove_last_textbox, width=3)
btn_remove.pack(side="left", padx=2)


# Crear y empacar el label "Pruebas"
pruebas_label = tk.Label(frame_izq, text="Pruebas", bg="white", font=FUENTE_GLOBAL)
pruebas_label.pack(anchor="w", padx=0, pady=2)

# Crear frame para los checkboxes
pruebas_frame = tk.Frame(frame_izq, bg="white")
pruebas_frame.pack(padx=0, pady=2)

# Crear checkboxes para cada prueba
for idx, prueba in enumerate(pruebas_opciones):
    var = tk.BooleanVar()
    checkbox = tk.Checkbutton(pruebas_frame, text=prueba, variable=var, bg="white", font=FUENTE_GLOBAL,
                              command=actualizar_vista)

    # Organizar checkboxes en dos columnas usando grid()
    row = idx // 2  # Determina la fila del checkbox
    col = idx % 2  # Determina la columna del checkbox (0 o 1)
    checkbox.grid(row=row, column=col, sticky="w", padx=10, pady=0)

    pruebas_checkboxes[prueba] = var

# Etiqueta para "Antiguedad"
antiguedad_label = tk.Label(frame_izq, text="Antiguedad", bg="white", font=FUENTE_GLOBAL)
antiguedad_label.pack(anchor="w", padx=0, pady=5)

# DateEntry para "Antiguedad"
antiguedad_date_entry = DateEntry(frame_izq, width=30, date_pattern='d/m/y', background="white", foreground="black", relief="groove", highlightthickness=1, font=FUENTE_GLOBAL)
antiguedad_date_entry.pack(fill=tk.X, padx=10, pady=0)

# Etiqueta para "Salario día"
salario_label = tk.Label(frame_izq, text="Salario día", bg="white", font=FUENTE_GLOBAL)
salario_label.pack(anchor="w", padx=0, pady=5)

# Entry para "Salario día"
salario_entry = tk.Entry(frame_izq, width=30, bg="white", fg="black", relief="sunken", highlightthickness=1, font=FUENTE_GLOBAL)
salario_entry.pack(fill=tk.X, padx=10, pady=0)


# Añadir Label encima del ScrolledText
label_retrasos = tk.Label(frame_izq, text="Relación de los retrasos", bg="white", font=FUENTE_GLOBAL)
label_retrasos.pack(anchor="w", padx=0, pady=5)

# Añadir ScrolledText widget
scrolled_text = ScrolledText(frame_izq, width=40, height=5, wrap=tk.WORD, bg="white", font=FUENTE_GLOBAL)
scrolled_text.pack(fill=tk.X, anchor="w", padx=10, pady=0)

# Botón para resolver
btn_resol = ttk.Button(frame_izq, text="Resolver", command=lambda: threading.Thread(target=thread_safe_resolver).start())
btn_resol.pack(anchor="e", pady=5, padx=10)


# Scrollbar
scrollbar = tk.Scrollbar(frame_der)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y, padx=1, pady=5)  # Agregar padding al scrollbar

# Text widget
vista_previa = tk.Text(frame_der, wrap=tk.WORD, yscrollcommand=scrollbar.set, bg="white", width=50, font=FUENTE_GLOBAL)
vista_previa.pack(fill=tk.BOTH, expand=True, padx=1, pady=5)  # Agregar padding al widget de texto

# Asociar el Scrollbar al Text widget
scrollbar.config(command=vista_previa.yview)

# Inicializar el Text widget con el contenido del modelo
vista_previa.insert(tk.END, modelo)


# Botón Guardar
btn_guardar = tk.Button(frame_izq, text="Guardar", command=combined_save_to_docx, bg="white", font=FUENTE_GLOBAL)
btn_guardar.pack(pady=0)

app.mainloop()
