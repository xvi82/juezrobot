import tkinter as tk
from tkinter import messagebox, ttk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkcalendar import DateEntry
from tkinter.scrolledtext import ScrolledText
import openai
import threading

# Configura tu clave de API
openai.api_key = 'CLAVE API AQUI'

# Leer las reglas para resolver el pleito
with open('resolutor3.txt', 'r', encoding="utf8") as file:
    reglas = file.read()


#crear hilo
def thread_safe_resolver():
    btn_resol.config(state=tk.DISABLED)  # Deshabilita el botón mientras se ejecuta la función
    mostrar_respuesta(año_var_1.get(), scrolled_text_1.get("1.0", tk.END).strip(), año_var_2.get(), scrolled_text_2.get("1.0", tk.END).strip())
    btn_resol.config(state=tk.NORMAL)  # Habilita el botón una vez que finaliza la función

# Obtener la respuesta de GPT-4 para resolver el pleito
def obtener_respuesta(prompt):
    # Hacer la petición a GPT-4
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": reglas},
            {"role": "user", "content": prompt},
        ]
    )
    texto = response['choices'][0]['message']['content'].strip()

    # Dividir el texto en frases (usando el punto como delimitador)
    frases = texto.split(".")

    # Inicializar variables
    respuesta_modificada = ""
    numero = None

    # Revisar cada frase para encontrar "CONDENAR" o "ABSOLVER"
    for frase in frases:
        if "CONDENAR" in frase:
            numero = 0
        elif "ABSOLVER" in frase:
            numero = 1
        else:
            respuesta_modificada += frase + "."

    # Limpieza final de la respuesta (eliminar el punto final)
    respuesta_modificada = respuesta_modificada.rstrip(".")

    return respuesta_modificada, numero

# Borra saltos de linea para GPT-4
def borrar_saltos_linea_vacios(texto):
        return texto.replace("\n\n", "\n")

# Mostrar la respuesta de GPT-4 en la GUI
# Recupera los valores
def mostrar_respuesta(año1, contenido1, año2, contenido2):
    global modelo  # Hacemos que modelo sea global para poder modificarla dentro de esta función
    respuesta, numero = obtener_respuesta("Las patologías y limitaciones de la parte demandante en el año " + año1 + " eran las siguientes: " + contenido1 + ".\n" + "Las patologías y limitaciones de la parte demandante en el año " + año2 + " eran las siguientes: " + contenido2 + ".")

    # Obtener el nombre del demandante y del demandado del cuadro de texto correspondiente
    nombre_demandante = entries["Nombre del demandante"].get()
    nombres_instituciones = ["INSS", "TGSS", "Mutua"]
    nombre_demandado_list = [nombre for nombre, var in instituciones_vars.items() if var.get()]
    if "Mutua" in nombre_demandado_list:
        mutua_nombre = nombre_mutua_entry.get()
        nombre_demandado_list.remove("Mutua")
        nombre_demandado_list.append(f"la Mutua {mutua_nombre}")

    if len(nombre_demandado_list) > 1:
        last_institucion = nombre_demandado_list.pop()
        nombre_demandado = ", ".join(nombre_demandado_list) + " y " + last_institucion
    else:
        nombre_demandado = ", ".join(nombre_demandado_list)
    tipo_incapacidad = var_incapacidad.get()
    # Reemplazar [argumentación por la IA] con la respuesta de GPT-4
    modelo_actualizado = modelo.replace("[argumentación por la IA]", borrar_saltos_linea_vacios(respuesta))
    print("Las patologías y limitaciones de la parte demandante en el año" + año1 + "eran las siguientes: " + contenido1 + ".\n" + "Las patologías y limitaciones de la parte demandante en el año" + año2 + "eran las siguientes: " + contenido2 + ".")
    # Reemplazar [fallo] según el número
    if numero == 0:
        fallo_texto = f"""
SEXTO.- En virtud de lo dispuesto en el art. 191.3 c) de la Ley Reguladora de la Jurisdicción Social, contra esta Sentencia puede interponerse Recurso de Suplicación

     Vistos los preceptos legales citados y demás de general observancia y por la autoridad que me confiere el art. 117 de la Constitución Española y 1 de la Ley Orgánica del Poder judicial,

FALLO

Que DEBO ESTIMAR y ESTIMO la demanda interpuesta por {nombre_demandante} frente {nombre_demandado}, reconociéndose al actor el grado de {tipo_incapacidad}, derivada de contingencia común, condenándose el Ente demandado a abonar al demandante de pensión, practicándose en ejecución de sentencia las compensaciones u opciones que procedan en atención a las prestaciones incompatibles que perciba.

Notifíquese la presente Resolución a las partes en legal forma, haciéndose saber al tiempo que contra la misma cabe recurso de Suplicación, para ante la Sala de lo Social del Tribunal Superior de Justicia de Canarias."""

    else:
        fallo_texto = f"""
SEXTO.- En virtud de lo dispuesto en el art. 191.3 c) de la Ley Reguladora de la Jurisdicción Social, contra esta Sentencia puede interponerse Recurso de Suplicación

     Vistos los preceptos legales citados y demás de general observancia y por la autoridad que me confiere el art. 117 de la Constitución Española y 1 de la Ley Orgánica del Poder judicial,

FALLO

Que DEBO DESESTIMAR y DESESTIMO la demanda interpuesta por {nombre_demandante} frente a {nombre_demandado}, y por ende absuelvo a la demanda de todos los pedimentos efectuados en su contra.

Notifíquese la presente Resolución a las partes en legal forma, haciéndose saber al tiempo que contra la misma cabe recurso de Suplicación, para ante la Sala de lo Social del Tribunal Superior de Justicia de Canarias."""

    modelo_actualizado = modelo_actualizado.replace("[fallo]", fallo_texto)

    vista_previa.delete(1.0, tk.END)
    vista_previa.insert(tk.END, modelo_actualizado)

    modelo = modelo_actualizado  # Actualizamos la variable global modelo con el contenido modificado



def cargar_modelo(filename):
    with open(filename, 'r', encoding='utf-8') as file:
        return file.read()


def formatear_fecha(date_obj):
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    dia = date_obj.day
    mes = meses[date_obj.month - 1]
    año = date_obj.year
    return f"{dia} de {mes} de {año}"


def actualizar_vista(event=None):
    modelo_temp = modelo[:]

    # Actualizar modelo_temp con las entradas de los widgets (Fecha, Número de juicio, etc.)
    for key, widget in entries.items():
        if key == "Fecha":
            fecha_seleccionada = widget.get_date()
            fecha_formateada = formatear_fecha(fecha_seleccionada)
            modelo_temp = modelo_temp.replace(f"[{key}]", fecha_formateada)
        else:
            modelo_temp = modelo_temp.replace(f"[{key}]", widget.get())

    # Construir los hechos probados escritos
    prefixes = ["PRIMERO", "SEGUNDO", "TERCERO", "CUARTO", "QUINTO", "SEXTO", "SÉPTIMO", "OCTAVO", "NOVENO", "DÉCIMO"]
    hechos_probados_escritos = "\n\n".join([prefixes[i] + ".- " + textbox.get() for i, textbox in enumerate(textboxes) if textbox.get()])

    # Determinar la continuación de la numeración para los hechos probados en el modelo
    num_hechos_probados_escritos = len([textbox for textbox in textboxes if textbox.get()])
    hechos_probados_modelo = modelo_temp.split("HECHOS PROBADOS\n")[1].split("FUNDAMENTOS DE DERECHO")[0].strip().split("\n")

    # Eliminamos la numeración anterior en el modelo
    hechos_probados_modelo_limpio = [(hecho.split(".-")[1] if ".-" in hecho else hecho) for hecho in hechos_probados_modelo if hecho.strip()]

    # Añadimos la nueva numeración
    hechos_probados_modelo_numerados = [prefixes[num_hechos_probados_escritos + i] + ".-" + hecho for i, hecho in enumerate(hechos_probados_modelo_limpio)]

    # Construir el texto final de hechos probados
    hechos_probados_final = hechos_probados_escritos + "\n\n" + "\n\n".join(hechos_probados_modelo_numerados)

    # Insertar los hechos probados final en el modelo
    modelo_temp = modelo_temp.split("HECHOS PROBADOS\n")[0] + "HECHOS PROBADOS\n\n" + hechos_probados_final + "\n\nFUNDAMENTOS DE DERECHO" + modelo_temp.split("FUNDAMENTOS DE DERECHO")[1]

    # Lógica para los checkboxes de pruebas
    pruebas_seleccionadas_list = [pruebas_mapping[prueba] for prueba, var in pruebas_checkboxes.items() if var.get()]
    if len(pruebas_seleccionadas_list) > 1:
        last_prueba = pruebas_seleccionadas_list.pop()
        pruebas_seleccionadas = ", ".join(pruebas_seleccionadas_list) + " y " + last_prueba
    else:
        pruebas_seleccionadas = ", ".join(pruebas_seleccionadas_list)

    modelo_temp = modelo_temp.replace("[pruebas]", pruebas_seleccionadas)

    # lógica seleccion de instituciones demandadas
    instituciones_seleccionadas = [institucion for institucion, var in instituciones_vars.items() if var.get()]
    if "Mutua" in instituciones_seleccionadas:
        nombre_mutua = nombre_mutua_entry.get().strip()
        if nombre_mutua:
            indice = instituciones_seleccionadas.index("Mutua")
            instituciones_seleccionadas[indice] = f"la Mutua '{nombre_mutua}'"
    if instituciones_seleccionadas:
        if len(instituciones_seleccionadas) > 1:
            instituciones_str = ", ".join(instituciones_seleccionadas[:-1]) + " y " + instituciones_seleccionadas[-1]
        else:
            instituciones_str = instituciones_seleccionadas[0]
        modelo_temp = modelo_temp.replace("[Nombre de la institución demandada]", instituciones_str)
    else:
        modelo_temp = modelo_temp.replace("[Nombre de la institución demandada]", "")

    vista_previa.delete(1.0, tk.END)
    vista_previa.insert(tk.END, modelo_temp)


def combined_save_to_docx():
    actualizar_vista()
    """Function to save the content from both the main GUI and the textboxes to a .docx file."""
    doc = Document()
    contenido = vista_previa.get(1.0, tk.END).split("\n")

    es_respuesta_inicio = False
    es_respuesta_final = False

    for linea in contenido:
        if "CUARTO.-" in linea:
            es_respuesta_inicio = True
            _agregar_linea_justificada_con_negrita(doc, linea)
        elif "QUINTO.-" in linea:
            es_respuesta_final = True
            _agregar_linea_justificada_con_negrita(doc, linea)
        elif es_respuesta_inicio and not es_respuesta_final:
            _agregar_linea_justificada(doc, linea)
        elif linea in ['SENTENCIA', 'ANTECEDENTES DE HECHO', 'FUNDAMENTOS DE DERECHO', 'FALLO']:
            _agregar_linea_centralizada_negrita(doc, linea)
        elif linea.startswith("TERCERO.-") and "autos se han observado" in linea:
            _agregar_linea_justificada_con_negrita(doc, linea)
        elif linea.startswith(('PRIMERO.-', 'SEGUNDO.-', 'TERCERO.-', 'CUARTO.-', 'QUINTO.-', 'SEXTO.-', 'SÉPTIMO', 'OCTAVO.-', 'NOVENO.-', 'DÉCIMO.-')):
            _agregar_linea_justificada_con_negrita(doc, linea)
        elif linea.startswith(("Visto por mí", 'Que DEBO', 'Vistos los preceptos legales', 'Notifíquese la presente', 'Así, por ésta, mi Sentencia')):
            _agregar_linea_justificada(doc, linea)
        elif linea == "HECHOS PROBADOS":
            _agregar_linea_centralizada_negrita(doc, linea)
        else:
            # Aquí es donde se agregan las líneas en blanco con el formato deseado
            par = doc.add_paragraph(linea)
            par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            par.paragraph_format.line_spacing = 1.0  # Espaciado entre líneas a 1,0
            par.paragraph_format.space_after = Pt(0)  # Espaciado posterior a 0 pt

    # El resto del código de la función para guardar el documento como .docx
    numero_procedimiento, año_procedimiento = entries["Número de juicio"].get().split("/")
    nombre_archivo = f"JO_revision_por_mejoria_{numero_procedimiento}-{año_procedimiento}.docx"
    doc.save(nombre_archivo)
    messagebox.showinfo("Guardado", f"Documento guardado como {nombre_archivo}")


def _agregar_linea_centralizada_negrita(doc, linea):
    par = doc.add_paragraph()
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = par.add_run(linea)
    run.bold = True
    par.paragraph_format.line_spacing = 1.0  # Espaciado entre líneas a 1,0
    par.paragraph_format.space_after = Pt(0)  # Espaciado posterior a 0 pt

def _agregar_linea_justificada_con_negrita(doc, linea):
    par = doc.add_paragraph()
    par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    start_word, rest_of_line = linea.split(" ", 1)
    run = par.add_run(start_word)
    run.bold = True
    par.add_run(f" {rest_of_line}")
    par.paragraph_format.line_spacing = 1.0  # Espaciado entre líneas a 1,0
    par.paragraph_format.space_after = Pt(0)  # Espaciado posterior a 0 pt


def _agregar_linea_justificada(doc, linea):
    par = doc.add_paragraph(linea)
    par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    par.paragraph_format.line_spacing = 1.0  # Espaciado entre líneas a 1,0
    par.paragraph_format.space_after = Pt(0)  # Espaciado posterior a 0 pt


def _crear_frame(parent):
    frame = tk.Frame(parent, relief="flat", borderwidth=1, bg="white")
    return frame

def add_textbox():
    """Función para añadir un nuevo cuadro de texto al UI."""
    textbox = tk.Entry(hechos_probados_frame, bg="white", fg="black", relief="sunken",
                       highlightthickness=1, font=FUENTE_GLOBAL)
    textbox.pack(fill=tk.X, pady=2, padx=10) # Asegurarse de que ocupe todo el ancho con fill=tk.X
    textbox.bind("<KeyRelease>", actualizar_vista)
    textboxes.append(textbox)

def remove_last_textbox():
    if textboxes:
        textbox = textboxes.pop()
        textbox.destroy()
        actualizar_vista()

def _crear_entradas(frame, keys):
    entries = {}
    for key in keys:
        label = tk.Label(frame, text=key, bg="white", font=FUENTE_GLOBAL)
        label.pack(anchor="w", padx=0, pady=0)

        if key == "Fecha":
            widget = DateEntry(frame, width=30, date_pattern='d/m/y', background="white", foreground="black",
                               relief="groove", highlightthickness=1)
            widget.configure(font=FUENTE_GLOBAL)
            widget.bind("<<DateEntrySelected>>", actualizar_vista)
        else:
            widget = tk.Entry(frame, width=30, bg="white", fg="black", relief="sunken", highlightthickness=1,
                              font=FUENTE_GLOBAL)
            widget.bind("<KeyRelease>", actualizar_vista)

        widget.pack(fill=tk.X, pady=2, padx=10)
        entries[key] = widget

    return entries

# Configuración inicial
modelo = cargar_modelo("revision mejoria.txt")  # Asegúrate de cambiar esta ruta al lugar correcto en tu sistema
keys = ['Fecha', 'Número de juicio', 'Nombre del demandante']

FUENTE_GLOBAL = ("Segoe UI", 9)  # Definición de la fuente global


# Configuración de la GUI
app = tk.Tk()
app.title("Rellenador de Modelo")
app.configure(bg="white")

# Establecer el tamaño de la ventana
app.state('zoomed')


# Centrar la ventana en la pantalla
screen_width = app.winfo_screenwidth()
screen_height = app.winfo_screenheight()
window_width = 1200
window_height = 750
x_position = (screen_width - window_width) / 2
y_position = (screen_height - window_height) / 2
app.geometry(f"{window_width}x{window_height}+{int(x_position)}+{int(y_position)}")

frame_izq = _crear_frame(app)
frame_der = _crear_frame(app)

# Diccionario para almacenar las variables de los checkboxes de pruebas
pruebas_checkboxes = {}

# Lista de pruebas disponibles
pruebas_opciones = ["Documental", "Testifical", "Perito actor", "Perito demandado", "Detective demandado", "Forense"]
pruebas_mapping = {
    "Documental": "la documental aportada",
    "Testifical": "la testifical",
    "Perito actor": "la pericial propuesta por la parte actora",
    "Perito demandado": "la pericial propuesta por la parte demandada",
    "Detective demandado": "la prueba de dectective propuesta por la parte demandada",
    "Forense": "la pericial forense",
}

# Configura la distribución de los frames usando grid
frame_izq.grid(row=0, column=0, sticky="nsew", padx=3, pady=3)  # Agregar padding a los frames
frame_der.grid(row=0, column=1, sticky="nsew", padx=3, pady=3)  # Agregar padding a los frames

# Configura cómo se distribuye el espacio entre las columnas
app.grid_columnconfigure(0, weight=1)  # 1/4 para el frame izquierdo
app.grid_columnconfigure(1, weight=5)  # 3/4 para el frame derecho
app.grid_rowconfigure(0, weight=1)

entries = _crear_entradas(frame_izq, keys)

instituciones_frame = _crear_frame(frame_izq)
instituciones_frame.pack(fill=tk.X, padx=0, pady=2)

instituciones_vars = {}
instituciones_opciones = ["INSS", "TGSS", "Mutua"]

label_instituciones = tk.Label(instituciones_frame, text="Nombre de la institución demandada", bg="white", font=FUENTE_GLOBAL)
label_instituciones.grid(row=0, column=0, sticky="w", padx=0, pady=2, columnspan=len(instituciones_opciones))

col_index = 0
for institucion in instituciones_opciones:
    var = tk.BooleanVar()
    checkbox = tk.Checkbutton(instituciones_frame, text=institucion, variable=var, bg="white", font=FUENTE_GLOBAL,
                              command=actualizar_vista)
    checkbox.grid(row=1, column=col_index, sticky="nsew", padx=10, pady=0)
    instituciones_vars[institucion] = var
    col_index += 1

# Hacer que cada columna se expanda equitativamente
for i in range(len(instituciones_opciones)):
    instituciones_frame.grid_columnconfigure(i, weight=1)

# Poner descripción de lo que hay que escribir
def handle_focus_in(_):
    if nombre_mutua_entry.get() == 'Nombre de la mutua...':
        nombre_mutua_entry.delete(0, tk.END)
        nombre_mutua_entry.config(fg='black')

def handle_focus_out(_):
    if not nombre_mutua_entry.get():
        nombre_mutua_entry.insert(0, 'Nombre de la mutua...')
        nombre_mutua_entry.config(fg='grey')

def handle_click(_):
    if nombre_mutua_entry.get() == 'Escriba el nombre de la mutua...':
        nombre_mutua_entry.delete(0, tk.END)
        nombre_mutua_entry.config(fg='black')


# Entry para el nombre de la mutua
nombre_mutua_entry = tk.Entry(instituciones_frame, bg="white", fg="black", relief="sunken", highlightthickness=1, font=FUENTE_GLOBAL)
nombre_mutua_entry.grid(row=3, column=0, sticky="e", padx=10, pady=2, columnspan=len(instituciones_opciones))

# Inicialmente, escondemos el label y el entry
nombre_mutua_entry.grid_remove()

nombre_mutua_entry.bind('<Button-1>', handle_click)


# Configura el Entry con el texto y el color predeterminados
nombre_mutua_entry.insert(0, 'Nombre de la mutua...')
nombre_mutua_entry.config(fg='grey')

# Vincula los eventos al Entry
nombre_mutua_entry.bind('<FocusIn>', handle_focus_in)
nombre_mutua_entry.bind('<FocusOut>', handle_focus_out)


def toggle_nombre_mutua(*args):
    if instituciones_vars["Mutua"].get():
        nombre_mutua_entry.grid()
    else:
        nombre_mutua_entry.grid_remove()
    actualizar_vista()

instituciones_vars["Mutua"].trace_add("write", toggle_nombre_mutua)
nombre_mutua_entry.bind("<KeyRelease>", lambda e: actualizar_vista())





# Add the label 'Hechos Probados' and the new textboxes to the existing GUI
hechos_probados_label = tk.Label(frame_izq, text="Hechos Probados", bg="white", font=FUENTE_GLOBAL)
hechos_probados_label.pack(anchor="w", padx=0, pady=2) # Añadido padding para consistencia

hechos_probados_frame = tk.Frame(frame_izq, bg="white")
hechos_probados_frame.pack(fill=tk.X, padx=0, pady=2) # Añadido fill=tk.X y padding para consistencia

# List to keep reference to all the textboxes
textboxes = []

# Add the first textbox
add_textbox()

# Frame para contener los botones
btn_frame = tk.Frame(frame_izq, bg="white")
btn_frame.pack(pady=5)

# Botón de adición
btn_add = ttk.Button(btn_frame, text="+", command=add_textbox, width=3)
btn_add.pack(side="left", padx=2)  # Se usa pack con side="left" para que estén en la misma línea

# Botón de eliminación
btn_remove = ttk.Button(btn_frame, text="-", command=remove_last_textbox, width=3)
btn_remove.pack(side="left", padx=2)

# Crear y empacar el label "Incapacidad Permanente revisada"
label_incapacidad = tk.Label(frame_izq, text="Incapacidad Permanente revisada", bg="white", font=FUENTE_GLOBAL)
label_incapacidad.pack(anchor="w", padx=0, pady=2)

# Definir las opciones para el combobox
opciones_incapacidad = ["Incapacidad Permanente Parcial",
                        "Incapacidad Permanente Total",
                        "Incapacidad Permanente Absoluta",
                        "Gran Invalidez"]

# Crear una variable tkinter StringVar para almacenar la opción seleccionada
var_incapacidad = tk.StringVar()

# Crear el combobox usando ttk.Combobox
combo_incapacidad = ttk.Combobox(frame_izq, textvariable=var_incapacidad, values=opciones_incapacidad, font=FUENTE_GLOBAL)
combo_incapacidad.pack(fill=tk.X, pady=2, padx=10)  # Asegurarse de que ocupe todo el ancho con fill=tk.X

# Crear y empacar el label "Pruebas"
pruebas_label = tk.Label(frame_izq, text="Pruebas", bg="white", font=FUENTE_GLOBAL)
pruebas_label.pack(anchor="w", padx=0, pady=2)

# Crear frame para los checkboxes
pruebas_frame = tk.Frame(frame_izq, bg="white")
pruebas_frame.pack(padx=0, pady=2)

# Crear checkboxes para cada prueba
for idx, prueba in enumerate(pruebas_opciones):
    var = tk.BooleanVar()
    checkbox = tk.Checkbutton(pruebas_frame, text=prueba, variable=var, bg="white", font=FUENTE_GLOBAL,
                              command=actualizar_vista)

    # Organizar checkboxes en dos columnas usando grid()
    row = idx // 2  # Determina la fila del checkbox
    col = idx % 2  # Determina la columna del checkbox (0 o 1)
    checkbox.grid(row=row, column=col, sticky="w", padx=10, pady=0)

    pruebas_checkboxes[prueba] = var

# Añadir Label encima del ScrolledText
label_retrasos = tk.Label(frame_izq, text="Patologías y limitaciones anteriores", bg="white", font=FUENTE_GLOBAL)
label_retrasos.pack(anchor="w", padx=0, pady=0)

# Frame para el primer Dropdown Menu
frame_año_1 = tk.Frame(frame_izq, bg="white")
frame_año_1.pack(fill=tk.X, padx=0, pady=5)

# Label y Dropdown Menu para el primer ScrolledText dentro del frame_año_1
label_año_1 = tk.Label(frame_año_1, text="Año:", bg="white", font=FUENTE_GLOBAL)
label_año_1.pack(side=tk.LEFT, padx=0, pady=0)

años = list(range(2023, 2009, -1))
año_var_1 = tk.StringVar()
año_var_1.set(años[0])  # valor por defecto
dropdown_1 = ttk.Combobox(frame_año_1, textvariable=año_var_1, values=años, font=FUENTE_GLOBAL, width=25)
dropdown_1.pack(fill=tk.X, anchor="w", padx=10, pady=0)


# Añadir ScrolledText widget
scrolled_text_1 = ScrolledText(frame_izq, width=40, height=4, wrap=tk.WORD, bg="white", font=FUENTE_GLOBAL)
scrolled_text_1.pack(fill=tk.X, anchor="w", padx=10, pady=0)

# Label y Dropdown Menu para el segundo ScrolledText
label_patologias = tk.Label(frame_izq, text="Patologías y limitaciones actuales", bg="white", font=FUENTE_GLOBAL)
label_patologias.pack(anchor="w", padx=0, pady=0)

# Frame para el segundo Dropdown Menu
frame_año_2 = tk.Frame(frame_izq, bg="white")
frame_año_2.pack(fill=tk.X, padx=0, pady=5)

# Label y Dropdown Menu para el segundo ScrolledText dentro del frame_año_2
label_año_2 = tk.Label(frame_año_2, text="Año:", bg="white", font=FUENTE_GLOBAL)
label_año_2.pack(side=tk.LEFT, padx=0, pady=0)

año_var_2 = tk.StringVar()
año_var_2.set(años[0])  # valor por defecto
dropdown_2 = ttk.Combobox(frame_año_2, textvariable=año_var_2, values=años, font=FUENTE_GLOBAL, width=25)
dropdown_2.pack(fill=tk.X, anchor="w", padx=10, pady=0)


# Segundo ScrolledText
scrolled_text_2 = ScrolledText(frame_izq, width=40, height=4, wrap=tk.WORD, bg="white", font=FUENTE_GLOBAL)
scrolled_text_2.pack(fill=tk.X, anchor="w", padx=10, pady=0)


# recuperar los datos
año_1 = año_var_1.get()
año_2 = año_var_2.get()
contenido_1 = scrolled_text_1.get("1.0", tk.END).strip()
contenido_2 = scrolled_text_2.get("1.0", tk.END).strip()


# Botón para resolver
btn_resol = ttk.Button(frame_izq, text="Resolver", command=lambda: threading.Thread(target=thread_safe_resolver).start())
btn_resol.pack(anchor="e", pady=5, padx=10)


# Scrollbar
scrollbar = tk.Scrollbar(frame_der)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y, padx=1, pady=5)  # Agregar padding al scrollbar

# Text widget
vista_previa = tk.Text(frame_der, wrap=tk.WORD, yscrollcommand=scrollbar.set, bg="white", width=50, font=FUENTE_GLOBAL)
vista_previa.pack(fill=tk.BOTH, expand=True, padx=1, pady=5)  # Agregar padding al widget de texto

# Asociar el Scrollbar al Text widget
scrollbar.config(command=vista_previa.yview)

# Inicializar el Text widget con el contenido del modelo
vista_previa.insert(tk.END, modelo)


# Botón Guardar
btn_guardar = tk.Button(frame_izq, text="Guardar", command=combined_save_to_docx, bg="white", font=FUENTE_GLOBAL)
btn_guardar.pack(pady=0)

app.mainloop()
